# services/word_service.py
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

TEMPLATE_PATH = "template.docx"
OUTPUT_DIR = "output"


def create_template():
    """建立有佔位符的 Word 範本（只需執行一次）"""
    doc = Document()
    doc.add_heading("客戶服務單", level=1)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    fields = [
        ("客戶姓名",         "{{name}}"),
        ("客戶聯絡方式(電話)", "{{phone}}"),
        ("想要的服務",        "{{service}}"),
        ("期望日期",          "{{date}}"),
        ("顏色喜好",          "{{color}}"),
        ("備註",             "{{note}}"),
    ]

    for i, (label, placeholder) in enumerate(fields):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = placeholder

    doc.save(TEMPLATE_PATH)
    print(f"✅ 範本已建立：{TEMPLATE_PATH}")


def fill_template(data: dict, session_id: str) -> str:
    """將蒐集到的資料填入 Word 範本，儲存為新檔案"""

    if not os.path.exists(TEMPLATE_PATH):
        create_template()

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document(TEMPLATE_PATH)

    # 替換表格中的佔位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value or "")

    # 用時間戳記命名，避免覆蓋
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{OUTPUT_DIR}/ticket_{session_id}_{timestamp}.docx"
    doc.save(filename)
    print(f"✅ 已輸出：{filename}")
    return filename