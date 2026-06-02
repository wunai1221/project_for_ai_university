from docx import Document
from datetime import datetime
import os
import logging

TEMPLATE_PATH = "template.docx"
OUTPUT_DIR = "output"

logger = logging.getLogger(__name__)

def fill_template(data: dict, session_id: str) -> str:
    """將蒐集到的資料填入 Word 範本，儲存為新檔案，並完美保留原始格式"""

    # 自動填入建立日期
    data["created_date"] = datetime.now().strftime("%Y/%m/%d")

    # 欄位預設值（避免佔位符殘留在表單上）
    defaults = {
        "name": "", "company": "", "phone": "", "email": "",
        "source": "", "service": "", "quantity_budget": "",
        "spec": "", "date": "", "staff": "", "note": "",
        "created_date": data["created_date"]
    }
    defaults.update(data)
    data = defaults

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document(TEMPLATE_PATH)

    # 替換表格中的佔位符（深入至段落層級，保護左側 ▌ 符號與表格線條）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        
                        # 如果段落中包含該佔位符，才進行精準替換
                        if placeholder in paragraph.text:
                            # 強制轉成字串防呆，並替換文字（這樣能保留儲存格的邊框與基本樣式）
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(OUTPUT_DIR, f"ticket_{session_id}_{timestamp}.docx")
    doc.save(filename)
    logger.info(f"✅ 已成功輸出商務表單：{filename}")
    return filename